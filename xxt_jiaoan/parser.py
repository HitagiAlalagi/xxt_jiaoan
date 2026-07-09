from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document

from .models import LessonItem
from .utils import clean_text




CHINESE_WEEK_NUMBERS = {
    "一": 1,
    "二": 2,
    "三": 3,
    "四": 4,
    "五": 5,
    "六": 6,
    "七": 7,
    "八": 8,
    "九": 9,
    "十": 10,
    "十一": 11,
    "十二": 12,
    "十三": 13,
    "十四": 14,
    "十五": 15,
    "十六": 16,
    "十七": 17,
    "十八": 18,
    "十九": 19,
    "二十": 20,
}


def all_doc_text(doc: Document) -> str:
    parts: list[str] = []
    for p in doc.paragraphs:
        text = clean_text(p.text)
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            row_text = [clean_text(cell.text) for cell in row.cells]
            parts.append("\n".join(t for t in row_text if t))
    return "\n".join(parts)


def table_pairs(doc: Document) -> dict[str, str]:
    pairs: dict[str, str] = {}
    for table in doc.tables:
        for row in table.rows:
            cells = unique_row_cells(row)
            if len(cells) >= 4 and cells[0] == "授课方式" and cells[2] == "课时安排":
                mode = re.sub(r"\s+", "", cells[1])
                if mode == "理论课":
                    pairs["theory_hours"] = cells[3]
                elif mode == "实训课":
                    pairs["practice_hours"] = cells[3]
                elif mode == "其他":
                    pairs["other_hours"] = cells[3]
                continue
            for i, cell in enumerate(cells):
                key = normalize_label(cell)
                if key and i + 1 < len(cells) and cells[i + 1]:
                    pairs[key] = cells[i + 1]
    return pairs


def unique_row_cells(row) -> list[str]:
    cells: list[str] = []
    last = None
    for cell in row.cells:
        text = clean_text(cell.text)
        if text != last:
            cells.append(text)
        last = text
    return cells


def template_fields(doc: Document) -> dict[str, str]:
    fields: dict[str, str] = {}
    label_map = {
        "授课题目": "topic",
        "学情分析": "analysis",
        "教学目的、要求": "objectives",
        "教学目标、要求": "objectives",
        "教学重点、难点及其解决方案": "key_points",
        "参考资料": "references",
        "教学反思": "reflection",
    }
    for table in doc.tables:
        for row in table.rows:
            row_text = "\n".join(unique_row_cells(row))
            normalized = clean_text(row_text)
            for label, key in label_map.items():
                if normalized.startswith(label):
                    value = re.sub(rf"^{re.escape(label)}\s*[:：]?\s*", "", normalized)
                    fields[key] = clean_text(value)
    return fields


def normalize_label(value: str) -> str:
    value = re.sub(r"[:：\s]", "", value or "")
    aliases = {
        "课程名称": "course_name",
        "授课时间": "date",
        "日期": "date",
        "开始节次": "start_period",
        "结束节次": "end_period",
        "课次": "lesson_no",
        "理论课": "theory_hours",
        "实训课": "practice_hours",
        "其他": "other_hours",
        "教材": "textbook",
        "教具": "teaching_aids",
        "授课题目": "topic",
        "学情分析": "analysis",
        "教学目标要求": "objectives",
        "教学目标、要求": "objectives",
        "教学目的要求": "objectives",
        "教学目的、要求": "objectives",
        "教学重点难点及其解决方案": "key_points",
        "教学重点、难点及其解决方案": "key_points",
        "参考资料": "references",
        "教学反思": "reflection",
        "课程类型": "course_type",
        "教学班级": "raw_class",
        "授课班级": "raw_class",
    }
    return aliases.get(value, "")


def find_after_label(text: str, labels: list[str], stop_labels: list[str]) -> str:
    label_re = "|".join(re.escape(x) for x in labels)
    stop_re = "|".join(re.escape(x) for x in stop_labels)
    m = re.search(rf"(?:{label_re})[:：]?\s*(.*?)(?=\n(?:{stop_re})[:：]?|\Z)", text, re.S)
    return clean_text(m.group(1)) if m else ""


def first_match(patterns: list[str], text: str) -> str:
    for pattern in patterns:
        m = re.search(pattern, text, re.S)
        if m:
            return clean_text(m.group(1))
    return ""


def normalize_date(value: str) -> str:
    value = clean_text(value)
    m = re.search(r"(\d{4})[-/.年](\d{1,2})[-/.月](\d{1,2})", value)
    if not m:
        return value
    return f"{int(m.group(1)):04d}-{int(m.group(2)):02d}-{int(m.group(3)):02d}"


def parse_schedule_time(value: str) -> tuple[int, str, str, str]:
    week = 0
    start = ""
    end = ""
    date = normalize_date(value)
    week_match = re.search(r"第\s*(\d{1,2})\s*周", value or "")
    period_match = re.search(r"第\s*(\d{1,2})\s*[~\-至]\s*(\d{1,2})\s*节", value or "")
    if week_match:
        week = int(week_match.group(1))
    if period_match:
        start = period_match.group(1)
        end = period_match.group(2)
    return week, start, end, date


def schedule_key(class_select: str, lesson_no: str) -> tuple[str, str]:
    return clean_text(class_select), clean_text(lesson_no)


def find_schedule_file(root: Path) -> Path | None:
    root = root.resolve()
    for parent in [root, *root.parents]:
        if parent.name.startswith("3."):
            course_root = parent.parent
            break
    else:
        course_root = root.parent
    for child in course_root.iterdir() if course_root.exists() else []:
        if child.is_dir() and child.name.startswith("1."):
            docs = [p for p in child.iterdir() if p.suffix.lower() == ".docx" and not p.name.startswith("~$")]
            if docs:
                return docs[0]
    return None


def class_names_from_items(items: list[LessonItem]) -> list[str]:
    names: list[str] = []
    for item in items:
        if item.class_select and item.class_select not in names:
            names.append(item.class_select)
    return sorted(names)


def score_schedule_for_class(rows: list[dict[str, str]], items: list[LessonItem], class_select: str) -> int:
    by_lesson = {
        clean_text(item.lesson_no): item
        for item in items
        if clean_text(item.class_select) == clean_text(class_select)
    }
    score = 0
    for row in rows:
        item = by_lesson.get(clean_text(row["lesson_no"]))
        if not item:
            continue
        if row["date"] and clean_text(item.date) == clean_text(row["date"]):
            score += 5
        if row["start_period"] and clean_text(item.start_period) == clean_text(row["start_period"]):
            score += 2
        if row["end_period"] and clean_text(item.end_period) == clean_text(row["end_period"]):
            score += 2
        if row["topic"] and clean_text(item.topic) == clean_text(row["topic"]):
            score += 3
    return score


def parse_schedule(root: Path, config: dict[str, Any], items: list[LessonItem] | None = None) -> dict[tuple[str, str], dict[str, str]]:
    configured_path = clean_text(str(config.get("schedule_file", "") or ""))
    path = Path(configured_path) if configured_path else None
    if path is None:
        found = find_schedule_file(root)
        if not found:
            return {}
        path = found
    doc = Document(str(path))
    configured_classes = config.get("schedule_class_order") or []
    schedule_tables = [
        table for table in doc.tables
        if len(table.rows) > 1 and "课次" in clean_text(table.rows[0].cells[0].text)
    ]
    table_rows: list[list[dict[str, str]]] = []
    for table in schedule_tables:
        rows: list[dict[str, str]] = []
        for row in table.rows[1:]:
            cells = unique_row_cells(row)
            if len(cells) < 3:
                continue
            lesson_no = clean_text(cells[0])
            if not lesson_no.isdigit():
                continue
            week, start, end, date = parse_schedule_time(cells[1])
            rows.append({
                "lesson_no": lesson_no,
                "week": str(week),
                "date": date,
                "start_period": start,
                "end_period": end,
                "topic": clean_text(cells[2]),
                "hours": clean_text(cells[4]) if len(cells) > 4 else "",
                "schedule_file": str(path),
            })
        if rows:
            table_rows.append(rows)

    configured_classes = [clean_text(x) for x in configured_classes if clean_text(x)]
    if configured_classes:
        class_names = configured_classes
    elif items:
        class_names = class_names_from_items(items)
    else:
        class_names = sorted(
            {
                clean_text(value)
                for value in config.get("class_map", {}).values()
                if clean_text(value)
            }
            or {
                p.name
                for p in root.rglob("*")
                if p.is_dir() and re.search(r".*班", p.name)
            }
        )

    assigned: dict[int, str] = {}
    remaining_classes = set(class_names)
    if items and class_names:
        scores = [
            (score_schedule_for_class(rows, items, class_select), table_index, class_select)
            for table_index, rows in enumerate(table_rows)
            for class_select in class_names
        ]
        for score, table_index, class_select in sorted(scores, reverse=True):
            if score <= 0 or table_index in assigned or class_select not in remaining_classes:
                continue
            assigned[table_index] = class_select
            remaining_classes.remove(class_select)
    for table_index, class_select in enumerate(class_names):
        if table_index < len(table_rows) and table_index not in assigned and class_select in remaining_classes:
            assigned[table_index] = class_select

    schedule: dict[tuple[str, str], dict[str, str]] = {}
    for table_index, rows in enumerate(table_rows):
        class_select = assigned.get(table_index)
        if not class_select:
            continue
        for row in rows:
            row = {**row, "class_select": class_select}
            schedule[schedule_key(class_select, row["lesson_no"])] = row
    return schedule


def align_items_with_schedule(items: list[LessonItem], root: Path, config: dict[str, Any]) -> list[str]:
    schedule = parse_schedule(root, config, items)
    if not schedule:
        return ["未找到授课计划，payload 未按授课计划校准。"]
    warnings: list[str] = []
    rows_by_class: dict[str, list[dict[str, str]]] = {}
    for row in schedule.values():
        rows_by_class.setdefault(row["class_select"], []).append(row)
    for class_select, class_items in group_items_by_class(items).items():
        unused_rows = rows_by_class.get(class_select, []).copy()
        for item in sorted(class_items, key=lambda x: (x.week, numeric_prefix(Path(x.file)), x.file)):
            row = best_schedule_row(item, unused_rows)
            if not row:
                warnings.append(f"{item.class_select} 第{item.lesson_no}次课未在授课计划中找到: {item.file}")
                continue
            unused_rows.remove(row)
            apply_schedule_row(item, row, warnings)
    return warnings


def group_items_by_class(items: list[LessonItem]) -> dict[str, list[LessonItem]]:
    grouped: dict[str, list[LessonItem]] = {}
    for item in items:
        grouped.setdefault(item.class_select, []).append(item)
    return grouped


def schedule_match_score(item: LessonItem, row: dict[str, str]) -> int:
    score = 0
    if clean_text(item.lesson_no) == clean_text(row["lesson_no"]):
        score += 8
    if item.date and clean_text(item.date) == clean_text(row["date"]):
        score += 10
    if item.start_period and clean_text(item.start_period) == clean_text(row["start_period"]):
        score += 3
    if item.end_period and clean_text(item.end_period) == clean_text(row["end_period"]):
        score += 3
    if item.week and str(item.week) == clean_text(row["week"]):
        score += 5
    item_topic = clean_text(item.topic)
    row_topic = clean_text(row["topic"])
    if item_topic and row_topic:
        if item_topic == row_topic:
            score += 10
        elif item_topic in row_topic or row_topic in item_topic:
            score += 4
    prefix = numeric_prefix(Path(item.file))
    if prefix and row["lesson_no"].isdigit():
        week_lesson_index = ((int(row["lesson_no"]) - 1) % 2) + 1
        if prefix == str(week_lesson_index):
            score += 2
    return score


def best_schedule_row(item: LessonItem, rows: list[dict[str, str]]) -> dict[str, str] | None:
    if not rows:
        return None
    scored = [(schedule_match_score(item, row), row) for row in rows]
    scored.sort(key=lambda x: x[0], reverse=True)
    if scored[0][0] <= 0:
        return None
    return scored[0][1]


def apply_schedule_row(item: LessonItem, row: dict[str, str], warnings: list[str]) -> None:
        before = {
            "date": item.date,
            "start_period": item.start_period,
            "end_period": item.end_period,
            "lesson_no": item.lesson_no,
            "topic": item.topic,
            "week": str(item.week),
        }
        item.date = row["date"] or item.date
        item.start_period = row["start_period"] or item.start_period
        item.end_period = row["end_period"] or item.end_period
        item.lesson_no = row["lesson_no"] or item.lesson_no
        item.topic = row["topic"] or item.topic
        if row["week"].isdigit():
            item.week = int(row["week"])
        for key, old_value in before.items():
            new_value = str(getattr(item, key))
            if old_value and new_value and clean_text(old_value) != clean_text(new_value):
                warnings.append(
                    f"已按授课计划修正 {item.class_select} 第{item.lesson_no}次课 {key}: {old_value} -> {new_value}"
                )


def parse_periods(value: str) -> tuple[str, str]:
    m = re.search(r"第\s*(\d{1,2})\s*节\s*至\s*第?\s*(\d{1,2})\s*节", value or "")
    if m:
        return m.group(1), m.group(2)
    return "", ""


def infer_week(path: Path, text: str) -> int:
    source = str(path) + "\n" + text[:500]
    m = re.search(r"第\s*(\d{1,2})\s*周", source)
    if m:
        return int(m.group(1))
    m = re.search(r"第\s*([一二三四五六七八九十]{1,3})\s*周", source)
    if m:
        return CHINESE_WEEK_NUMBERS.get(m.group(1), 0)
    return 0


def infer_kind(path: Path, text: str) -> str:
    parent_name = path.parent.name
    path_text = str(path)
    for source in (parent_name, path.name):
        m = re.search(r"((?:计算机应用技术|大数据技术|口腔修复工艺|医疗设备安装与维护)\d{4}班?)", source)
        if m:
            return m.group(1)
    if parent_name and not re.match(r"第\d+周", parent_name):
        return parent_name
    m = re.search(r"(计算机应用技术\d{4}班)", path_text)
    if m:
        return m.group(1)
    if "医疗" in parent_name:
        return "医疗"
    if "口腔" in parent_name:
        return "口腔"
    if "医疗" in text[:500]:
        return "医疗"
    return ""


def class_select_for(raw_class: str, kind: str, config: dict[str, Any]) -> str:
    class_map = config.get("class_map", {})
    if kind in class_map:
        return class_map[kind]
    if raw_class in class_map:
        return class_map[raw_class]
    normalized = clean_text(raw_class)
    normalized = normalized.replace("（五年制）班", "（五年制）")
    normalized = normalized.replace("(五年制)班", "(五年制)")
    return normalized or kind


def numeric_prefix(path: Path) -> str:
    m = re.match(r"(\d+)[_、.\-]", path.name)
    return m.group(1) if m else ""


def find_continuation_file(docx_path: Path, kind: str) -> str:
    candidates = [
        p for p in docx_path.parent.glob("*.docx")
        if not p.name.startswith("~$") and "续页" in p.name
    ]
    prefix = numeric_prefix(docx_path)
    if prefix:
        for p in candidates:
            if f"续页-{prefix}" in p.stem:
                return str(p)
    preferred = f"续页-1" if kind == "口腔" else f"续页-2" if kind == "医疗" else ""
    for p in candidates:
        if preferred and preferred in p.stem:
            return str(p)
    return str(candidates[0]) if candidates else ""


def parse_docx(path: Path, config: dict[str, Any]) -> LessonItem:
    doc = Document(str(path))
    text = all_doc_text(doc)
    pairs = table_pairs(doc)
    fields = template_fields(doc)
    kind = infer_kind(path, text)
    week = infer_week(path, text)

    stop = [
        "教学目标、要求",
        "教学目标",
        "教学重点、难点及其解决方案",
        "参考资料",
        "教学反思",
        "教案续页",
        "授课题目",
        "学情分析",
    ]

    date = normalize_date(pairs.get("date", "") or first_match([
        r"授课时间[:：]?\s*([0-9年月日./-]+)",
        r"日期[:：]?\s*([0-9年月日./-]+)",
    ], text))
    parsed_start, parsed_end = parse_periods(pairs.get("date", ""))
    raw_class = pairs.get("raw_class", "")

    item = LessonItem(
        file=str(path),
        course_name=pairs.get("course_name") or config.get("course_query", ""),
        course_type=pairs.get("course_type") or config.get("default_course_type", ""),
        raw_class=raw_class,
        date=date,
        start_period=pairs.get("start_period") or parsed_start or config.get("period_map", {}).get(kind, {}).get("start_period", ""),
        end_period=pairs.get("end_period") or parsed_end or config.get("period_map", {}).get(kind, {}).get("end_period", ""),
        lesson_no=pairs.get("lesson_no") or str(week),
        theory_hours=pairs.get("theory_hours") or config.get("hours", {}).get("theory_hours", ""),
        practice_hours=pairs.get("practice_hours") or config.get("hours", {}).get("practice_hours", ""),
        other_hours=pairs.get("other_hours") or config.get("hours", {}).get("other_hours", ""),
        textbook=pairs.get("textbook") or config.get("textbook", ""),
        teaching_aids=pairs.get("teaching_aids") or config.get("teaching_aids", ""),
        topic=fields.get("topic") or pairs.get("topic") or find_after_label(text, ["授课题目"], ["学情分析"]),
        analysis=fields.get("analysis") or pairs.get("analysis") or find_after_label(text, ["学情分析"], ["教学目标、要求", "教学目标", "教学目的、要求"]),
        objectives=fields.get("objectives") or pairs.get("objectives") or find_after_label(text, ["教学目的、要求", "教学目标、要求", "教学目标"], stop),
        key_points=fields.get("key_points") or pairs.get("key_points") or find_after_label(text, ["教学重点、难点及其解决方案"], stop),
        references=fields.get("references") or pairs.get("references") or find_after_label(text, ["参考资料"], stop),
        reflection=fields.get("reflection") or pairs.get("reflection") or find_after_label(text, ["教学反思"], ["教案续页"]),
        week=week,
        kind=kind,
        class_select=class_select_for(raw_class, kind, config),
        continuation_file=find_continuation_file(path, kind),
    )
    return item


def discover_docx(root: Path) -> list[Path]:
    return sorted(
        p for p in root.rglob("*.docx")
        if not p.name.startswith("~$")
        and "续页" not in p.name
        and "模版" not in p.name
        and "模板" not in p.name
    )
